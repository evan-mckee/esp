"""
ESP Resume Maker 2.0
7/12/2021
Evan McKee
"""


import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EspApp:
    """
    Evan-Style-Python or ESP:
    Decouple the input variables and data from your code and store them in an external JSON file.
    Read the file into a dict at the beginning; update during runtime, save and load whenever.

    Resume Maker:
    Formats and saves a resume as word docx using a JSON file as input.
    """

    def __init__(self, esp_data, output_file):
        """
        Read the JSON file and initialize runtime variables.

        :param esp_data: Path to JSON file to read (str)
        :param output_file: Path to JSON file to write (str)
        """
        
        # Read in the JSON data / input variables
        self.d = self.read_esp_data(esp_data)

        self.output_file = output_file
        self.doc = Document()
        self.section = self.doc.sections[0]
        margins = self.d['parameters']['margins_tbl']
        self.section.top_margin = Inches(margins[0])
        self.section.bottom_margin = Inches(margins[1])
        self.section.left_margin = Inches(margins[2])
        self.section.right_margin = Inches(margins[3])

    def run(self):
        """
        Write the marked up resume using 'resume' and 'format' dicts.
        """

        for line in self.d['resume']:
            [indent, fmat, txt] = line.split('|')
            p = self.doc.add_paragraph()
            p_f = p.paragraph_format
            style = {}
            for entry in ['Font', "Size", "Bold", "Special", "Color", "Alignment", "Bullets", "Underline"]:
                if entry in self.d['formats'][fmat]:
                    style[entry] = self.d['formats'][fmat][entry]
                else:
                    style[entry] = self.d['formats']['Text'][entry]  # 'Text' style is default
            if style['Bullets'] is True:
                p.style = 'List Bullet'
            if style['Alignment'] == 'Left':
                p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Bring to center by splitting into n lines
            # Add text
            thisrun = p.add_run(txt)
            if style['Special'] == 'All Caps':
                thisrun.font.all_caps = True
            elif style['Special'] == 'Small Caps':
                thisrun.font.small_caps = True
            else:
                thisrun.font.all_caps = False
                thisrun.font.small_caps = False
            p_f.left_indent = Inches(int(indent) * self.d['parameters']['cascade_indent'])
            thisrun.font.name = style['Font']
            thisrun.font.size = Pt(style['Size'])
            c = style['Color']
            thisrun.font.color.rgb = RGBColor(c, c, c)
            thisrun.font.bold = style['Bold']
            thisrun.font.underline = style['Underline']
            p_f.space_after = Pt(self.d['parameters']['vertical_spacing'])

        self.doc.save(self.d['parameters']['word_output'])

    def read_esp_data(self, data_file):
        """
        Read JSON data into a dict.

        :param data_file: Path to JSON file (str)
        :return: Dict of data (dict)
        """
        with open(data_file, 'r') as f:
            return json.load(f)

    def write_esp_data(self, data_file):
        """
        Output modified dict data into JSON file.

        :param data_file: Target json file (str)
        """
        with open(data_file, 'w') as f:
            json.dump(self.d, f, indent=4, sort_keys=True)


A = EspApp('esp_resume_data.json', 'esp_resume_data.json')
A.run()
